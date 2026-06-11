"""Testes da extração multi-formato de documentos de cliente (content_library)."""
from __future__ import annotations

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest  # noqa: E402

from core.services.content_library import extract_document_text  # noqa: E402


def _make_docx(paragraphs, table=None) -> bytes:
    import docx
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_paragrafos_e_tabela():
    data = _make_docx(["Proposta ACME", "Robôs autônomos."],
                      table=[["Rubrica", "Valor"], ["Pessoal", "100k"]])
    text = extract_document_text(data, "proposta.docx")
    assert "Proposta ACME" in text
    assert "Robôs autônomos." in text
    assert "Rubrica | Valor" in text  # tabela serializada
    assert "Pessoal | 100k" in text


def test_txt_e_md():
    assert extract_document_text(b"linha um\nlinha dois", "notas.txt") == "linha um\nlinha dois"
    assert "# Titulo" in extract_document_text(b"# Titulo\n\ncorpo", "doc.md")


def test_formato_nao_suportado():
    with pytest.raises(RuntimeError, match="não suportado"):
        extract_document_text(b"\x00\x01", "antigo.doc")
    with pytest.raises(RuntimeError, match="não suportado"):
        extract_document_text(b"...", "planilha.xlsx")
