"""
Content Library — operações CRUD e extração de fatos atômicos via LLM.

Cada ContentItem pertence a um workspace. Os campos `summary`, `key_facts`
e `themes` são gerados por LLM no momento do upload (ou manualmente via /enrich).
"""
from __future__ import annotations

import json
import logging
import os
import re
import uuid
from io import BytesIO

from supabase import Client

logger = logging.getLogger(__name__)

CONTENT_TYPES = {"proposal", "project_description", "team_bio", "technical_doc", "other"}

_ENRICH_SYSTEM = """Você é um especialista em análise de documentos de inovação e P&D.
Extraia informações estruturadas do documento abaixo. Responda APENAS com JSON válido."""

_ENRICH_USER = """Documento:
\"\"\"
{content}
\"\"\"

Extraia e retorne:
{{
  "summary": "Resumo em 3 frases do que este documento trata.",
  "key_facts": [
    "Fato atômico 1 relevante para uma proposta de P&D",
    "Fato atômico 2",
    "..."
  ],
  "themes": ["tema1", "tema2", "tema3"],
  "importance_score": 1-10
}}

key_facts deve ter entre 5 e 15 itens. Cada fato deve ser autocontido (não usar pronomes como 'eles', 'a empresa').
themes deve ter entre 2 e 6 termos técnicos ou setoriais.
importance_score reflete o valor do documento para escrita de propostas de fomento:
  1-3 = trivial/genérico (sobre nós, missão genérica)
  4-6 = útil como material de apoio (descrição padrão, blog post)
  7-9 = rico e específico (projetos concretos, resultados quantificados, propostas aprovadas)
  10 = referência crítica (proposta aprovada com edital similar, contrato relevante)"""


# =============================================================================
# LLM
# =============================================================================

def _make_client():
    from core.llm.llm_client import make_client
    backend = os.getenv("LLM_BACKEND", "openai").lower()
    if backend == "gemini":
        return make_client(
            api_key=os.environ["GEMINI_API_KEY"],
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
        ), "gemini-2.5-flash"
    return make_client(api_key=os.environ["OPENAI_API_KEY"]), os.getenv("OPENAI_MODEL", "gpt-4o-mini")


def enrich_content(content: str) -> dict:
    """Chama LLM para extrair summary, key_facts e themes de um texto."""
    client, model = _make_client()
    truncated = content[:8000]  # evitar tokens excessivos
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": _ENRICH_SYSTEM},
                {"role": "user", "content": _ENRICH_USER.format(content=truncated)},
            ],
            temperature=0.1,
            max_tokens=1500,
        )
        raw = response.choices[0].message.content.strip()
        if "```" in raw:
            raw = re.sub(r"```(?:json)?", "", raw).strip()
        data = json.loads(raw)
        # importance_score: aceita int 1-10; fallback para 5 (default neutro)
        raw_score = data.get("importance_score", 5)
        try:
            importance = max(1, min(10, int(raw_score)))
        except (TypeError, ValueError):
            importance = 5
        return {
            "summary": data.get("summary", ""),
            "key_facts": data.get("key_facts", []),
            "themes": data.get("themes", []),
            "importance_score": importance,
        }
    except Exception as e:
        logger.error("Erro ao enriquecer conteúdo: %s", e)
        return {"summary": "", "key_facts": [], "themes": [], "importance_score": 5}


# =============================================================================
# PDF EXTRACTION
# =============================================================================

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extrai texto de um PDF via pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(BytesIO(pdf_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in pages if p.strip())
    except ImportError:
        raise RuntimeError("pypdf não instalado. Execute: pip install pypdf") from None
    except Exception as e:
        raise RuntimeError(f"Falha ao extrair texto do PDF: {e}") from e


def _extract_text_from_docx(data: bytes) -> str:
    """Extrai texto de um .docx via python-docx (parágrafos + tabelas).

    .docx é o formato mais comum de proposta de cliente. Leve (sem torch) —
    roda no caminho síncrono do upload sem penalizar latência. .doc binário
    antigo NÃO é suportado pelo python-docx (orientar a salvar como .docx).
    """
    try:
        import docx  # python-docx
        doc = docx.Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except ImportError:
        raise RuntimeError("python-docx não instalado. Execute: pip install python-docx") from None
    except Exception as e:
        raise RuntimeError(f"Falha ao extrair texto do DOCX: {e}") from e


# Formatos aceitos no upload da library (docs do cliente → perfil/escrita).
SUPPORTED_UPLOAD_EXTS = ("pdf", "docx", "txt", "md")


def extract_document_text(data: bytes, filename: str) -> str:
    """Extrai texto de um documento de cliente, roteando por extensão.

    Multi-formato (docs de cliente são heterogêneos): PDF, DOCX, TXT, MD.
    Parsers leves no caminho síncrono do upload — Docling (estruturado, pesado)
    fica reservado ao chunking em background, onde o ganho de layout se aplica.
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""
    if ext == "pdf":
        return extract_text_from_pdf(data)
    if ext == "docx":
        return _extract_text_from_docx(data)
    if ext in ("txt", "md"):
        return data.decode("utf-8", errors="ignore")
    raise RuntimeError(
        f"Formato .{ext or '?'} não suportado. Use: {', '.join(SUPPORTED_UPLOAD_EXTS)} "
        "(.doc antigo → salve como .docx)."
    )


# =============================================================================
# CRUD
# =============================================================================

def get_workspace_id(db: Client, user_id: str) -> str:
    """Retorna ou cria workspace para o user_id.

    Recebe o cliente Supabase do caller para que a operação respeite RLS
    (cliente autenticado com JWT do request em rotas FastAPI).
    """
    result = db.table("workspaces").select("id").eq("user_id", user_id).limit(1).execute()
    if result.data:
        return result.data[0]["id"]
    created = db.table("workspaces").insert({"user_id": user_id, "profile": {}}).execute()
    return created.data[0]["id"]


def list_items(
    db: Client, workspace_id: str, include_archived: bool = False
) -> list[dict]:
    """Lista items do workspace.

    Por padrão filtra fora rows soft-deleted (archived_at IS NOT NULL).
    Passe include_archived=True para uma "archive view" que retorna todos.
    """
    q = (
        db.table("content_items")
        .select("id, title, type, tags, summary, themes, enrichment_status, created_at, updated_at, archived_at")
        .eq("workspace_id", workspace_id)
    )
    if not include_archived:
        q = q.is_("archived_at", "null")
    result = q.order("created_at", desc=True).execute()
    return result.data or []


def get_item(db: Client, item_id: str, workspace_id: str) -> dict | None:
    result = (
        db.table("content_items")
        .select("*")
        .eq("id", item_id)
        .eq("workspace_id", workspace_id)
        .maybe_single()
        .execute()
    )
    return result.data if result else None


async def create_item(
    db: Client,
    workspace_id: str,
    title: str,
    type_: str,
    content: str,
    tags: list[str],
    source_url: str | None = None,
    enrich: bool = True,
) -> dict:
    """Insere um content_item e (opcionalmente) agenda enriquecimento via LLM.

    O INSERT é sempre síncrono e imediato — a linha é criada com
    summary/key_facts/themes vazios. Se `enrich=True` (padrão), um job
    procrastinate é enfileirado para preencher esses campos em background.

    Callers que não querem enfileirar o job (testes, importações em massa)
    devem passar `enrich=False`.
    """
    item_id = str(uuid.uuid4())
    result = db.table("content_items").insert({
        "id": item_id,
        "workspace_id": workspace_id,
        "title": title,
        "type": type_,
        "content": content,
        "source_url": source_url,
        "tags": tags,
        "summary": "",
        "key_facts": [],
        "themes": [],
    }).execute()
    item = result.data[0]

    if enrich:
        # Import local para evitar ciclo: core.tasks importa core.services.content_library.
        from core.tasks import app as tasks_app

        async with tasks_app.open_async():
            await tasks_app.configure_task("enrich_content").defer_async(
                item_id=str(item["id"])
            )

    return item


async def update_item(
    db: Client, item_id: str, workspace_id: str, updates: dict
) -> dict | None:
    """Atualiza item da library.

    Quando `content` muda: zera summary/key_facts/themes/embedding, recoloca
    importance_score=5 (default neutro), marca enrichment_status='pending' e
    enfileira `enrich_content_task` (que encadeia `embed_content_task`). Mesma
    chain do `create_item` — evita inconsistência entre summary/themes novos e
    embedding stale do conteúdo anterior (Gap 1 / ADR M8).

    Quando só metadados (title/type/tags/source_url) mudam: UPDATE direto, sem
    LLM, sem queue.
    """
    allowed = {"title", "type", "content", "tags", "source_url"}
    payload = {k: v for k, v in updates.items() if k in allowed}
    if not payload:
        return get_item(db, item_id, workspace_id)

    content_changed = False
    if "content" in payload:
        current = get_item(db, item_id, workspace_id)
        if current is None:
            return None
        if (current.get("content") or "") != payload["content"]:
            content_changed = True

    if content_changed:
        # Zera enriched para não servir summary/themes/embedding do conteúdo
        # antigo enquanto o worker reprocessa. importance volta ao neutro
        # (default do schema) para não enviesar o retrieval multi-critério.
        payload["summary"] = ""
        payload["key_facts"] = []
        payload["themes"] = []
        payload["embedding"] = None
        payload["importance_score"] = 5
        payload["enrichment_status"] = "pending"

    result = (
        db.table("content_items")
        .update(payload)
        .eq("id", item_id)
        .eq("workspace_id", workspace_id)
        .execute()
    )
    if not result.data:
        return None

    if content_changed:
        # Import local para evitar ciclo: core.tasks importa core.services.content_library.
        from core.tasks import app as tasks_app

        async with tasks_app.open_async():
            await tasks_app.configure_task("enrich_content").defer_async(
                item_id=str(item_id)
            )

    return result.data[0]


def delete_item(db: Client, item_id: str, workspace_id: str) -> bool:
    result = (
        db.table("content_items")
        .delete()
        .eq("id", item_id)
        .eq("workspace_id", workspace_id)
        .execute()
    )
    return bool(result.data)


def mark_items_referenced(
    db: Client, item_ids: list[str], workspace_id: str
) -> None:
    """Atualiza last_referenced_at para uma lista de items (ADR B4).

    Chamada quando items são injetados em uma WritingSession (via construtor
    com `library_items` ou via @ mention em um turno). O timestamp alimenta
    a fórmula de decay temporal: effective_importance = importance_score *
    exp(-(now - last_referenced_at) / half_life). Falhas são silenciadas com
    log para não interromper o fluxo principal.
    """
    if not item_ids:
        return
    try:
        from datetime import datetime, timezone
        now_iso = datetime.now(timezone.utc).isoformat()
        (
            db.table("content_items")
            .update({"last_referenced_at": now_iso})
            .in_("id", item_ids)
            .eq("workspace_id", workspace_id)
            .execute()
        )
    except Exception as e:
        logger.warning("mark_items_referenced falhou para %d items: %s", len(item_ids), e)


def search_items(
    db: Client, workspace_id: str, query: str, type_filter: str | None = None
) -> list[dict]:
    """Busca simples por ilike no título, summary e tags.

    Exclui items arquivados (archived_at IS NOT NULL) por padrão.
    """
    q = (
        db.table("content_items")
        .select("id, title, type, tags, summary, themes, created_at")
        .eq("workspace_id", workspace_id)
        .is_("archived_at", "null")
        .ilike("title", f"%{query}%")
    )
    if type_filter:
        q = q.eq("type", type_filter)
    result = q.order("created_at", desc=True).execute()
    return result.data or []


def archive_item(db: Client, item_id: str, workspace_id: str) -> bool:
    """Soft-delete: marca archived_at = now() se ainda não estiver arquivado.

    Returns True se uma row foi atualizada, False caso contrário (não existe,
    pertence a outro workspace, ou já estava arquivada).
    """
    result = (
        db.table("content_items")
        .update({"archived_at": "now()"})
        .eq("id", item_id)
        .eq("workspace_id", workspace_id)
        .is_("archived_at", "null")
        .execute()
    )
    return bool(result.data)


def unarchive_item(db: Client, item_id: str, workspace_id: str) -> bool:
    """Reverte soft-delete: zera archived_at."""
    result = (
        db.table("content_items")
        .update({"archived_at": None})
        .eq("id", item_id)
        .eq("workspace_id", workspace_id)
        .execute()
    )
    return bool(result.data)
